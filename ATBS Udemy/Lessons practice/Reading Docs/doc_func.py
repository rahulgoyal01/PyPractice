from docx import Document
import docx
d = docx.Document('D:\\demo.docx')
d.paragraphs[0]
d.paragraphs[0].text
d.paragraphs[1].text
d.paragraphs[2].text

p = d.paragraphs[2]
p.runs
p.runs[0].text
p.runs[1].text
p.runs[2].text
p.runs[3].text
p.runs[4].text

p.runs[2].bold = True
p.runs[4].italic = True
p.runs[2].underline = True

d.save('D:\\demo2.docx')

'''Document object have paragraph object and para 
object has run object to add more lines to paragraph
and modify it'''

#p.style
#p.style = 'Title'
#d.save('D:\\demo2.docx')

